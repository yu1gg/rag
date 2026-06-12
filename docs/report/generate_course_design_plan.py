from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import BaseStyle


DOC_TITLE = "深度学习应用实践课程设计方案"
PROJECT_TITLE = "基于检索增强生成（RAG）的智能问答与文本摘要系统设计与实现"
COURSE_NAME = "深度学习应用实践"
DEFAULT_STUDENT_NAME = "【学生姓名待填】"
DEFAULT_CLASS_NAME = "【班级待填】"
DEFAULT_ADVISOR = "【指导教师待填】"
DEFAULT_OUTPUT = Path(__file__).with_name("project_X_实践课程设计方案.docx")

COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_BLUE = RGBColor(0x2E, 0x74, 0xB5)
COLOR_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_LIGHT_FILL = "F2F4F7"
COLOR_BORDER = "C8D2DC"


@dataclass(frozen=True)
class ParagraphBlock:
    text: str


@dataclass(frozen=True)
class BulletBlock:
    items: Sequence[str]
    numbered: bool = False


@dataclass(frozen=True)
class TableBlock:
    headers: Sequence[str]
    rows: Sequence[Sequence[str]]
    column_widths: Sequence[float]


@dataclass(frozen=True)
class Section:
    title: str
    blocks: Sequence[object]


SECTIONS: Sequence[Section] = (
    Section(
        title="一、选题背景",
        blocks=(
            ParagraphBlock(
                "随着大语言模型在知识问答、文本理解与内容生成任务中的广泛应用，"
                "如何让模型在限定知识范围内稳定、可信地回答问题，已经成为课程实践中具有代表性的研究方向。"
                "单纯依赖通用大模型进行开放式回答，容易出现事实性偏差、引用不清晰和回答脱离资料的问题。"
            ),
            ParagraphBlock(
                "检索增强生成（RAG）通过“先检索、后生成”的方式，将外部知识库与大语言模型结合起来，"
                "既保留了语言生成的自然性，又增强了回答的依据性和可追溯性。"
                "同时，长文本摘要也是教学资料整理、学习辅助和知识提炼中的常见需求，适合与问答能力共同构成一套综合性实践系统。"
            ),
            ParagraphBlock(
                "基于以上背景，本课题拟围绕 project_X 项目，设计并实现一个集数据采集、文本清洗、知识检索、智能问答、"
                "文本摘要和效果评估于一体的课程实践方案，并以当前 Ver 2.0 的 FastAPI + Vue 工程骨架作为实现载体，"
                "提升方案的工程完整性与可落地性。"
            ),
        ),
    ),
    Section(
        title="二、设计目标",
        blocks=(
            ParagraphBlock(
                "本课题计划构建一个基于 RAG 架构的智能问答与文本摘要系统。系统面向课程资料、公开论文摘要、新闻文本和自建问答对，"
                "实现知识组织、检索增强回答与长文本摘要等功能。"
            ),
            BulletBlock(
                items=(
                    "完成多源文本数据的收集、整理与统一存储，形成可用于课程实践的基础知识库。",
                    "建立文本清洗、去重、分块和索引构建流程，为后续检索与生成提供高质量输入。",
                    "基于 BGE-M3、FAISS 与 reranker 构建检索链路，实现与大语言模型协同的问答与摘要功能。",
                    "通过 FastAPI 提供标准化后端接口，并结合 Vue 前端形成可交互的课程实践系统。",
                    "使用自动指标与人工评价相结合的方法，对检索质量、回答质量与摘要质量进行综合评估。",
                ),
                numbered=True,
            ),
        ),
    ),
    Section(
        title="三、数据收集",
        blocks=(
            ParagraphBlock(
                "本课题的数据来源采用“公开资料 + 本地课程资料 + 自建问答对”的组合方式，以兼顾知识覆盖度、实验可控性和项目可重复性。"
                "所有原始数据统一保存为 JSON 结构，并输出到后端存储目录中，便于后续清洗和索引构建。"
            ),
            TableBlock(
                headers=("数据来源", "采集方式", "计划用途"),
                rows=(
                    ("arXiv 论文摘要", "调用 arXiv 接口或现成包采集", "补充前沿概念与技术背景"),
                    ("课程资料", "读取本地 Markdown、TXT、PDF 等文本", "构建课程场景下的核心知识库"),
                    ("公开新闻与技术资讯", "通过 requests 获取公开可访问资料", "扩展案例与应用场景"),
                    ("自建 QA 对", "结合课程主题手工编写或模板生成", "补充高频问答与评估样本"),
                ),
                column_widths=(1.45, 2.15, 2.90),
            ),
            ParagraphBlock(
                "数据收集阶段计划达到不少于 50 篇文档、总字符量不少于 10 万字，同时补充不少于 200 组自建问答对。"
                "若在线采集阶段出现网络波动或内容不足，可使用本地课程资料和模板化补充文档作为兜底，以保证课程设计流程能够继续推进。"
            ),
        ),
    ),
    Section(
        title="四、数据预处理",
        blocks=(
            ParagraphBlock(
                "为保证检索与生成效果，本课题将在数据进入索引前完成统一的文本预处理。"
                "预处理不仅关注文本格式规范化，还要尽量减少重复、噪声和过短片段对检索结果的干扰。"
            ),
            BulletBlock(
                items=(
                    "文本清洗：去除 HTML 标签、异常空白符和无意义控制字符，统一 UTF-8 编码格式。",
                    "重复过滤：依据文本指纹或近似去重策略过滤高相似文档，避免知识库中同质片段过多。",
                    "短文本筛除：对过短、信息量不足的样本进行过滤，减少低质量片段进入索引。",
                    "滑动分块：按照固定长度与重叠窗口进行切分，优先考虑段落边界，保证语义连续性。",
                    "结构化输出：分别生成 `documents.json`、`qa_pairs.json`、`chunks.json` 等中间产物，便于后续脚本复用。",
                ),
                numbered=True,
            ),
            ParagraphBlock(
                "在工程实现上，预处理流程将对接 `collect_data`、`clean_data` 与 `chunk_data` 等批处理脚本，"
                "并将运行期产物统一落入 `backend/storage/raw` 与 `backend/storage/processed` 目录。"
            ),
        ),
    ),
    Section(
        title="五、模型设计",
        blocks=(
            ParagraphBlock(
                "本课题采用“嵌入检索 + 重排序 + 大模型生成”的 RAG 总体架构。"
                "系统先对文档块和用户查询进行向量化，再在向量索引中完成相似度检索，随后利用重排序模型筛选最相关片段，"
                "最后将检索上下文注入提示词模板，交由大语言模型生成最终回答或摘要。"
            ),
            BulletBlock(
                items=(
                    "嵌入模型：采用 `BAAI/bge-m3` 对文本块与用户查询进行向量表示。",
                    "向量索引：采用 `FAISS` 保存向量及元数据，实现高效 Top-K 检索。",
                    "重排序模型：采用 `bge-reranker-v2-m3` 对候选片段进行二次排序，提高相关性。",
                    "生成模型：采用 OpenAI 兼容接口调用大语言模型，支持问答模式与摘要模式。",
                    "服务封装：通过 `RagService` 统一封装索引加载、检索、问答、摘要与引用结果组织逻辑。",
                ),
            ),
            ParagraphBlock(
                "与模板示例中的单一文本分类任务不同，本课题更强调知识组织与生成式应用的协同。"
                "因此模型设计不仅包括算法链路本身，还包括后端接口设计、前端交互入口与评估闭环，形成更接近真实应用场景的实践系统。"
            ),
        ),
    ),
    Section(
        title="六、实验方案",
        blocks=(
            ParagraphBlock(
                "实验方案采用“离线构建知识库、在线完成问答与摘要、评估阶段统一分析结果”的方式组织。"
                "离线阶段主要完成数据准备和索引构建，在线阶段通过接口与前端页面验证系统的可交互性，最后结合自动评估与人工评价分析方案效果。"
            ),
            TableBlock(
                headers=("实验环节", "主要配置或实现方案"),
                rows=(
                    ("数据准备", "执行 `collect_data`、`clean_data`、`chunk_data`，生成原始数据与文本块"),
                    ("索引构建", "执行 `build_index`，输出 `faiss.index` 与 `chunk_meta.json`"),
                    ("检索生成", "BGE-M3 向量检索 + reranker 重排 + OpenAI 兼容 LLM"),
                    ("后端实现", "使用 FastAPI 提供 `/rag/qa`、`/rag/summary`、`/rag/search` 等接口"),
                    ("前端验证", "使用 Vue 3 + Vite 提供聊天式或表单式交互页面，辅助人工测试"),
                ),
                column_widths=(1.40, 5.10),
            ),
            BulletBlock(
                items=(
                    "运行数据采集、清洗、分块与建索引脚本，完成离线知识库准备。",
                    "使用健康检查与状态接口确认数据量、索引状态和基础服务可用性。",
                    "在问答模式下提交课程相关问题，观察回答内容与引用片段是否一致。",
                    "在摘要模式下输入长文本，检查摘要是否满足结构化与压缩要求。",
                    "整理若干测试样本，分别进行自动评分与人工打分，形成评估记录。",
                ),
                numbered=True,
            ),
        ),
    ),
    Section(
        title="七、评价方法",
        blocks=(
            ParagraphBlock(
                "本课题的评价方法采用自动指标与人工评价并行的方式。"
                "自动指标用于观察检索命中与生成文本的整体质量，人工评价用于补充对正确性、完整性和可读性的主观判断。"
            ),
            TableBlock(
                headers=("评价维度", "评价指标或方法", "说明"),
                rows=(
                    ("检索效果", "Hit@K", "判断正确参考片段是否出现在检索结果中"),
                    ("摘要质量", "ROUGE、BLEU", "衡量生成摘要与参考摘要之间的重叠程度"),
                    ("语义相似度", "BERTScore", "评估生成文本与参考答案的语义接近程度"),
                    ("人工评价", "正确性、完整性、流畅性", "采用 5 分制记录主观质量并分析典型案例"),
                ),
                column_widths=(1.25, 1.95, 3.30),
            ),
            ParagraphBlock(
                "在人工评价阶段，可选取 20 组左右具有代表性的测试样本，分别记录成功案例、边界案例与失败案例，"
                "分析其问题成因是否来自数据质量、检索偏差、提示词设计或模型生成能力不足。"
            ),
        ),
    ),
    Section(
        title="八、预期成果",
        blocks=(
            ParagraphBlock(
                "本课程设计预期形成从数据准备到系统实现、再到效果评估的完整实践成果，主要包括以下内容："
            ),
            BulletBlock(
                items=(
                    "一套围绕 RAG 问答与文本摘要任务的数据采集、清洗、分块与索引构建流程。",
                    "一个基于 BGE-M3、FAISS、reranker 与大语言模型的检索增强生成系统原型。",
                    "基于 FastAPI 的后端服务接口，以及基于 Vue 的前端交互页面或工作台。",
                    "一组可复用的离线批处理脚本，包括数据准备、索引构建与评估脚本。",
                    "自动指标报告、人工评价记录和典型案例分析材料。",
                    "课程设计方案、课程设计报告以及后续答辩展示所需的文档基础。",
                ),
                numbered=True,
            ),
        ),
    ),
    Section(
        title="九、可行性分析",
        blocks=(
            ParagraphBlock(
                "从技术可行性来看，本课题所使用的关键组件均已有较成熟的开源实现，包括 BGE 系列嵌入模型、FAISS 向量检索、"
                "OpenAI 兼容接口调用方式以及 Python 与 Vue 的常见工程化框架，因此整体技术路线明确、复现成本较低。"
            ),
            ParagraphBlock(
                "从工程可行性来看，当前仓库已经具备 Ver 2.0 的前后端分层基础，后端可通过脚本完成数据准备与索引构建，"
                "前端可承担结果展示与交互验证，能够为课程设计提供完整的实现载体，而不必从零搭建全部基础设施。"
            ),
            ParagraphBlock(
                "从实验可行性来看，本课题的数据来源可控，评估指标明确，且可根据硬件条件灵活调整采样规模、Top-K 参数、"
                "模型调用频率与测试样本数量。若外部接口或在线采集受限，还可使用本地资料与模板化内容进行兜底，因此具备按时完成课程设计的现实条件。"
            ),
        ),
    ),
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成 project_X 实践课程设计方案 DOCX。")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="输出 DOCX 路径")
    parser.add_argument("--student-name", default=DEFAULT_STUDENT_NAME, help="学生姓名")
    parser.add_argument("--class-name", default=DEFAULT_CLASS_NAME, help="班级")
    parser.add_argument("--advisor", default=DEFAULT_ADVISOR, help="指导教师")
    parser.add_argument("--project-title", default=PROJECT_TITLE, help="实践课题标题")
    return parser.parse_args()


def set_style_font(style: BaseStyle, font_name: str, east_asia_name: str, size: int, color: RGBColor, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def format_run(run, *, font_name: str = "Calibri", east_asia_name: str = "微软雅黑", size: int = 11, color: RGBColor = COLOR_BLACK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, *, color: str = COLOR_BORDER, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_tag = tc_borders.find(qn(f"w:{edge}"))
        if edge_tag is None:
            edge_tag = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_tag)
        edge_tag.set(qn("w:val"), "single")
        edge_tag.set(qn("w:sz"), str(size))
        edge_tag.set(qn("w:color"), color)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        child = tc_mar.find(qn(f"w:{key}"))
        if child is None:
            child = OxmlElement(f"w:{key}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_document(document: Document) -> None:
    for section in document.sections:
        configure_page(section)

    normal_style = document.styles["Normal"]
    set_style_font(normal_style, "Calibri", "微软雅黑", 11, COLOR_BLACK)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.1
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    heading1 = document.styles["Heading 1"]
    set_style_font(heading1, "Calibri", "微软雅黑", 16, COLOR_BLUE, bold=True)
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.line_spacing = 1.1

    heading2 = document.styles["Heading 2"]
    set_style_font(heading2, "Calibri", "微软雅黑", 13, COLOR_BLUE, bold=True)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing = 1.1

    heading3 = document.styles["Heading 3"]
    set_style_font(heading3, "Calibri", "微软雅黑", 12, COLOR_DARK_BLUE, bold=True)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)
    heading3.paragraph_format.line_spacing = 1.1

    list_bullet = document.styles["List Bullet"]
    set_style_font(list_bullet, "Calibri", "微软雅黑", 11, COLOR_BLACK)
    list_bullet.paragraph_format.space_after = Pt(8)
    list_bullet.paragraph_format.line_spacing = 1.167

    list_number = document.styles["List Number"]
    set_style_font(list_number, "Calibri", "微软雅黑", 11, COLOR_BLACK)
    list_number.paragraph_format.space_after = Pt(8)
    list_number.paragraph_format.line_spacing = 1.167


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.paragraph_format.space_after = Pt(0)
    header_run = header_para.add_run("project_X | 实践课程设计方案")
    format_run(header_run, size=9, color=COLOR_MUTED)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_after = Pt(0)
    footer_run = footer_para.add_run("深度学习应用实践")
    format_run(footer_run, size=9, color=COLOR_MUTED)


def add_spacer(document: Document, after_pt: int) -> None:
    para = document.add_paragraph()
    para.paragraph_format.space_after = Pt(after_pt)


def add_cover(document: Document, student_name: str, class_name: str, advisor: str, project_title: str) -> None:
    add_spacer(document, 16)

    course_para = document.add_paragraph()
    course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_para.paragraph_format.space_after = Pt(8)
    run = course_para.add_run(COURSE_NAME)
    format_run(run, size=12, color=COLOR_MUTED, bold=True)

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(DOC_TITLE)
    format_run(run, size=23, color=COLOR_BLACK, bold=True)

    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(20)
    run = subtitle_para.add_run(project_title)
    format_run(run, size=14, color=COLOR_MUTED)

    meta_table = document.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    widths = (Inches(3.10), Inches(3.10))
    labels = (
        ("学生姓名", student_name),
        ("班级", class_name),
        ("指导教师", advisor),
        ("课程名称", COURSE_NAME),
    )

    cell_pairs = (
        (meta_table.cell(0, 0), labels[0]),
        (meta_table.cell(0, 1), labels[1]),
        (meta_table.cell(1, 0), labels[2]),
        (meta_table.cell(1, 1), labels[3]),
    )
    for cell, (label, value) in cell_pairs:
        cell.width = widths[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        set_cell_border(cell)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = para.add_run(f"{label}：")
        format_run(label_run, size=10.5, color=COLOR_DARK_BLUE, bold=True)
        value_run = para.add_run(value)
        format_run(value_run, size=10.5)

    merged_cell = meta_table.cell(2, 0).merge(meta_table.cell(2, 1))
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(merged_cell)
    set_cell_border(merged_cell)
    para = merged_cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    label_run = para.add_run("实践课题：")
    format_run(label_run, size=10.5, color=COLOR_DARK_BLUE, bold=True)
    value_run = para.add_run(project_title)
    format_run(value_run, size=10.5)

    for row in meta_table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                cell.width = widths[index]

    add_spacer(document, 14)
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_after = Pt(18)
    intro_run = intro.add_run(
        "本方案围绕 RAG 智能问答与文本摘要系统的设计思路、实验流程和预期成果展开，"
        "用于指导课程实践阶段的任务实施。"
    )
    format_run(intro_run, size=10.5, color=COLOR_MUTED, italic=True)


def add_section(document: Document, section: Section) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading_run = heading.add_run(section.title)
    format_run(heading_run, size=16, color=COLOR_BLUE, bold=True)

    for block in section.blocks:
        if isinstance(block, ParagraphBlock):
            para = document.add_paragraph(style="Normal")
            para.paragraph_format.first_line_indent = Inches(0.28)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(block.text)
        elif isinstance(block, BulletBlock):
            style_name = "List Number" if block.numbered else "List Bullet"
            for item in block.items:
                para = document.add_paragraph(style=style_name)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(item)
        elif isinstance(block, TableBlock):
            table = document.add_table(rows=1, cols=len(block.headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            for index, header in enumerate(block.headers):
                cell = table.rows[0].cells[index]
                cell.width = Inches(block.column_widths[index])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                set_cell_border(cell)
                set_cell_shading(cell, COLOR_LIGHT_FILL)
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(header)
                format_run(run, size=10.5, color=COLOR_DARK_BLUE, bold=True)

            for row_data in block.rows:
                row = table.add_row()
                for index, value in enumerate(row_data):
                    cell = row.cells[index]
                    cell.width = Inches(block.column_widths[index])
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_margins(cell)
                    set_cell_border(cell)
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_after = Pt(0)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = para.add_run(value)
                    format_run(run, size=10.5)
            document.add_paragraph()


def build_document(student_name: str, class_name: str, advisor: str, project_title: str) -> Document:
    document = Document()
    configure_document(document)
    add_header_footer(document)
    add_cover(document, student_name, class_name, advisor, project_title)

    for section in SECTIONS:
        add_section(document, section)

    return document


def iter_document_text(document: Document) -> Iterable[str]:
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            yield paragraph.text.strip()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield paragraph.text.strip()


def validate_document(path: Path, student_name: str, class_name: str, advisor: str, project_title: str) -> None:
    document = Document(path)
    text_blocks = list(iter_document_text(document))
    full_text = "\n".join(text_blocks)

    assert DOC_TITLE in full_text, "缺少文档主标题"
    assert project_title in full_text, "缺少实践课题标题"
    expected_titles = [section.title for section in SECTIONS]
    for title in expected_titles:
        assert title in full_text, f"缺少章节：{title}"
    assert len(document.tables) == 4, f"表格数量不符合预期：{len(document.tables)}"
    assert "BERT预训练模型的外卖评论情感分类" not in full_text, "模板原始课题内容未被完全替换"
    for keyword in ("RAG", "FastAPI", "Vue"):
        assert keyword in full_text, f"缺少关键术语：{keyword}"

    placeholder_count = full_text.count(DEFAULT_STUDENT_NAME) + full_text.count(DEFAULT_CLASS_NAME) + full_text.count(DEFAULT_ADVISOR)
    if student_name == DEFAULT_STUDENT_NAME and class_name == DEFAULT_CLASS_NAME and advisor == DEFAULT_ADVISOR:
        assert placeholder_count == 3, f"占位符数量异常：{placeholder_count}"
    else:
        assert placeholder_count == 0, "自定义封面信息时仍检测到默认占位符"


def save_document(document: Document, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    args = parse_args()
    document = build_document(
        student_name=args.student_name,
        class_name=args.class_name,
        advisor=args.advisor,
        project_title=args.project_title,
    )
    save_document(document, args.output)
    validate_document(
        args.output,
        student_name=args.student_name,
        class_name=args.class_name,
        advisor=args.advisor,
        project_title=args.project_title,
    )
    print(f"Document generated: {args.output}")


if __name__ == "__main__":
    main()
